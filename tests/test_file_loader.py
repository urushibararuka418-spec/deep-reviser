"""
FileLoader 单元测试
===================
测试多格式文档加载器：txt (UTF-8/GBK/GB2312/UTF-16)、docx、epub，
以及不支持格式的异常处理。
"""

import os
import sys
import tempfile
import pytest

# 确保项目 src 目录在 sys.path 中
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from preprocessor.file_loader import FileLoader


# ── 辅助工具 ──────────────────────────────────────────────

def _write_temp_file(suffix: str, content: bytes | str, encoding: str | None = None) -> str:
    """创建临时文件并返回路径。"""
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    if isinstance(content, str):
        with open(path, "w", encoding=encoding or "utf-8") as f:
            f.write(content)
    else:
        with open(path, "wb") as f:
            f.write(content)
    return path


# ── 测试：txt 文件加载 ────────────────────────────────────

class TestLoadTxt:
    """测试纯文本文件的加载，覆盖多种编码。"""

    def test_load_txt_utf8(self):
        """加载 UTF-8 编码的 txt 文件。"""
        content = "Hello World\n你好，世界！\n"
        path = _write_temp_file(".txt", content, encoding="utf-8")
        try:
            loader = FileLoader()
            result = loader.load(path)
            assert result == content
            assert "Hello World" in result
            assert "你好，世界" in result
        finally:
            os.unlink(path)

    def test_load_txt_gbk(self):
        """加载 GBK 编码的 txt 文件（中文 Windows 常见编码）。"""
        content = "GBK编码测试\n简体中文内容\n"
        path = _write_temp_file(".txt", content.encode("gbk"))
        try:
            loader = FileLoader()
            result = loader.load(path)
            assert "GBK编码测试" in result
            assert "简体中文内容" in result
        finally:
            os.unlink(path)

    def test_load_txt_gb2312(self):
        """加载 GB2312 编码的 txt 文件。"""
        content = "GB2312编码测试\n"
        path = _write_temp_file(".txt", content.encode("gb2312"))
        try:
            loader = FileLoader()
            result = loader.load(path)
            assert "GB2312编码测试" in result
        finally:
            os.unlink(path)

    def test_load_txt_utf16(self):
        """加载 UTF-16 编码的 txt 文件（含 BOM）。"""
        content = "UTF-16 编码测试\n"
        path = _write_temp_file(".txt", content.encode("utf-16"))
        try:
            loader = FileLoader()
            result = loader.load(path)
            assert "UTF-16 编码测试" in result
        finally:
            os.unlink(path)

    def test_load_md_file(self):
        """加载 .md Markdown 文件（应走文本加载逻辑）。"""
        content = "# 标题\n\n正文内容\n"
        path = _write_temp_file(".md", content, encoding="utf-8")
        try:
            loader = FileLoader()
            result = loader.load(path)
            assert result == content
            assert "# 标题" in result
        finally:
            os.unlink(path)


# ── 测试：不支持格式 ──────────────────────────────────────

class TestUnsupportedFormat:
    """测试不支持的文件格式应抛出 ValueError。"""

    def test_load_unsupported_pdf(self):
        """.pdf 不是支持格式，应抛出 ValueError。"""
        # 创建一个假的 pdf 文件
        path = _write_temp_file(".pdf", b"%PDF-1.4 fake")
        try:
            loader = FileLoader()
            with pytest.raises(ValueError, match="不支持的文件格式"):
                loader.load(path)
        finally:
            os.unlink(path)

    def test_load_no_extension(self):
        """无后缀名文件应抛出 ValueError。"""
        path = _write_temp_file("", b"no extension")
        try:
            loader = FileLoader()
            with pytest.raises(ValueError, match="不支持的文件格式"):
                loader.load(path)
        finally:
            os.unlink(path)

    def test_load_unsupported_binary(self):
        """.exe 不是支持格式，应抛出 ValueError。"""
        path = _write_temp_file(".exe", b"\x4d\x5a")
        try:
            loader = FileLoader()
            with pytest.raises(ValueError, match="不支持的文件格式"):
                loader.load(path)
        finally:
            os.unlink(path)


# ── 测试：docx 文件加载 ───────────────────────────────────

class TestLoadDocx:
    """测试 Word 文档 (.docx) 的加载。"""

    def test_load_docx(self):
        """加载简单的 .docx 文件，提取段落文本。"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        # 创建临时 docx 文件
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc = Document()
            doc.add_paragraph("第一段内容")
            doc.add_paragraph("第二段：包含中文和 English。")
            doc.add_heading("章节标题", level=1)
            doc.add_paragraph("标题后的段落。")
            doc.save(path)

            loader = FileLoader()
            result = loader.load(path)
            assert "第一段内容" in result
            assert "第二段：包含中文和 English。" in result
            assert "章节标题" in result
            assert "标题后的段落" in result
        finally:
            os.unlink(path)

    def test_load_empty_docx(self):
        """加载空白的 .docx 文件。"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc = Document()
            doc.save(path)

            loader = FileLoader()
            result = loader.load(path)
            # 空白文档应返回空字符串
            assert result == ""
        finally:
            os.unlink(path)


# ── 测试：epub 文件加载 ───────────────────────────────────

class TestLoadEpub:
    """测试 EPUB 电子书的加载。"""

    def test_load_epub(self):
        """加载简单的 .epub 文件，提取文本内容。"""
        try:
            from ebooklib import epub
            from bs4 import BeautifulSoup
        except ImportError:
            pytest.skip("ebooklib 或 beautifulsoup4 未安装")

        # 创建一个极简的 epub 文件
        fd, path = tempfile.mkstemp(suffix=".epub")
        os.close(fd)
        try:
            book = epub.EpubBook()
            book.set_identifier("test-001")
            book.set_title("测试电子书")
            book.set_language("zh")

            # 创建章节
            c1 = epub.EpubHtml(
                title="第一章",
                file_name="chap_01.xhtml",
                lang="zh",
            )
            c1.content = "<h1>第一章 标题</h1><p>这是第一章的内容。</p>"

            c2 = epub.EpubHtml(
                title="第二章",
                file_name="chap_02.xhtml",
                lang="zh",
            )
            c2.content = "<p>第二章内容，包含 <strong>加粗</strong> 文字。</p>"

            book.add_item(c1)
            book.add_item(c2)
            book.spine = ["nav", c1, c2]

            # 添加必要的导航和toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            epub.write_epub(path, book)

            loader = FileLoader()
            result = loader.load(path)
            assert "第一章 标题" in result
            assert "这是第一章的内容" in result
            assert "第二章内容" in result
            assert "加粗" in result
        finally:
            os.unlink(path)

    def test_supported_formats(self):
        """验证 SUPPORTED 常量包含正确的格式。"""
        assert ".txt" in FileLoader.SUPPORTED
        assert ".md" in FileLoader.SUPPORTED
        assert ".docx" in FileLoader.SUPPORTED
        assert ".epub" in FileLoader.SUPPORTED
        assert ".pdf" not in FileLoader.SUPPORTED
        assert ".exe" not in FileLoader.SUPPORTED


# ── 测试：文件不存在 ──────────────────────────────────────

class TestFileNotFound:
    """测试加载不存在的文件应抛出 FileNotFoundError。"""

    def test_load_nonexistent_file(self):
        """加载不存在的文件路径。"""
        loader = FileLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("/tmp/nonexistent_file_xyz123.txt")
